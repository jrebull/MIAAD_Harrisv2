"""Genera un .docx con Título + Abstract + Keywords, extraídos del .tex de envío."""
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

TEX = Path("/Users/haowei/Documents/MIAAD/SMART/Harris2/MICAI/main_compacta_submission.tex")
OUT = Path("/Users/haowei/Documents/MIAAD/SMART/Harris2/MICAI/Two-Condition-Diagnostic_title-abstract.docx")
src = TEX.read_text()


def clean(s):
    s = re.sub(r"\\texorpdfstring\{\\\\\}\{ \}", " ", s)
    s = re.sub(r"\\emph\{([^}]*)\}", r"\1", s)
    s = re.sub(r"\\textbf\{([^}]*)\}", r"\1", s)
    s = s.replace("{,}", ",")
    s = s.replace("---", "—").replace("--", "–")
    s = re.sub(r"\\ ", " ", s)        # U.S.\  -> U.S.
    s = re.sub(r"\\,", " ", s)   # thin space
    s = re.sub(r"\\%", "%", s)
    s = re.sub(r"\\&", "&", s)
    s = re.sub(r"\\cite\{[^}]*\}", "", s)
    s = re.sub(r"\s+", " ", s).strip()
    return s


title = clean(re.search(r"\\title\{(.+?)\}\n", src, re.S).group(1))
abstract = clean(re.search(r"\\begin\{abstract\}(.+?)\\keywords", src, re.S).group(1))
kw_raw = re.search(r"\\keywords\{(.+?)\}\n", src, re.S).group(1)
keywords = ", ".join(clean(k) for k in kw_raw.split(r"\and"))

doc = Document()
style = doc.styles["Normal"]
style.font.name = "Times New Roman"
style.font.size = Pt(11)

t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run(title)
r.bold = True
r.font.size = Pt(15)

doc.add_paragraph()  # espacio

h = doc.add_paragraph()
h.add_run("Abstract").bold = True
ap = doc.add_paragraph(abstract)
ap.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

doc.add_paragraph()
k = doc.add_paragraph()
k.add_run("Keywords: ").bold = True
k.add_run(keywords)

doc.save(OUT)
words = len(abstract.split())
print(f"OK -> {OUT}")
print(f"título: {title}")
print(f"abstract: {words} palabras")
print(f"keywords: {keywords}")
